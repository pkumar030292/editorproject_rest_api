import re
import os
import uuid
from docx import Document
from docx.table import _Cell
from typing import List, Dict
from datetime import datetime

PLACEHOLDER_PATTERN = r"\{\{\{(.*?)\}\}\}"

def extract_placeholders(template_path: str) -> List[str]:
    """
    Extract all triple-brace placeholders like {{{label}}} from a .docx file.
    Includes paragraphs, table cells, headers, and footers.
    """
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")

    doc = Document(template_path)
    found = set()

    # Check paragraphs
    for para in doc.paragraphs:
        matches = re.findall(PLACEHOLDER_PATTERN, para.text)
        found.update(matches)

    # Check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                matches = re.findall(PLACEHOLDER_PATTERN, cell.text)
                found.update(matches)

    # Check headers and footers
    for section in doc.sections:
        header = section.header
        footer = section.footer
        for para in header.paragraphs + footer.paragraphs:
            matches = re.findall(PLACEHOLDER_PATTERN, para.text)
            found.update(matches)

    return list(found)


def _replace_placeholders_in_paragraph(para, context: Dict[str, str]):
    for key, value in context.items():
        placeholder = f"{{{{{{{key}}}}}}}"
        if placeholder in para.text:
            para.text = para.text.replace(placeholder, value)

def _replace_placeholders_in_paragraph(para, context: dict):
    for key, value in context.items():
        placeholder = f"{{{{{{{key}}}}}}}"
        if placeholder in para.text:
            para.text = para.text.replace(placeholder, value)



def generate_hld_doc(template_path: str, context: dict, output_dir: str = "output") -> str:
    """
    Replace placeholders in .docx template with values from context.
    Filename format: <template_name>_YYYYMMDD_HHMMSS.docx
    Returns the generated filename.
    """
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")

    doc = Document(template_path)

    # Replace in paragraphs
    for para in doc.paragraphs:
        _replace_placeholders_in_paragraph(para, context)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _replace_placeholders_in_paragraph(cell, context)

    # Replace in headers and footers
    for section in doc.sections:
        for para in section.header.paragraphs + section.footer.paragraphs:
            _replace_placeholders_in_paragraph(para, context)

    os.makedirs(output_dir, exist_ok=True)

    # Generate filename
    template_name = os.path.splitext(os.path.basename(template_path))[0]
    timestamp = datetime.now().strftime("%Y%m%d_%I-%M-%S%p")
    filename = f"{template_name}_{timestamp}.docx"

    output_path = os.path.join(output_dir, filename)
    doc.save(output_path)
    return filename

