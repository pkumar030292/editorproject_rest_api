import uuid
from pathlib import Path
from docx2pdf import convert
from fastapi import FastAPI, HTTPException, Query, Request, Form
from fastapi.responses import FileResponse, JSONResponse,HTMLResponse
from pydantic import BaseModel, Field
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from typing import Dict
from utils.doc_generator import extract_placeholders, generate_hld_doc
import os
import logging
from io import BytesIO
from fastapi.responses import JSONResponse
from docx import Document
from utils.doc_generator import extract_placeholders, generate_hld_doc, _replace_placeholders_in_paragraph



# Configure logging
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")

app = FastAPI(title="Dynamic HLD Generator")

# Directories
TEMPLATE_DIR = "templates"
OUTPUT_DIR = "output"
# Directories
BASE_DIR = Path(__file__).resolve().parent
TEMPLATE_DIR = BASE_DIR / "templates"
OUTPUT_DIR = BASE_DIR / "outputs"
OUTPUT_DIR.mkdir(exist_ok=True)

# Serve static files (optional for JS/CSS)
app.mount("/static", StaticFiles(directory="static"), name="static")
# Setup Jinja2 templates
templates = Jinja2Templates(directory="templates")

@app.get("/", response_class=HTMLResponse)
def home_page(request: Request):
    templates_list = [t for t in os.listdir(TEMPLATE_DIR) if t.endswith(".docx")]
    files_list = [f for f in os.listdir(OUTPUT_DIR) if f.endswith(".docx")]

    return templates.TemplateResponse("index.html", {
        "request": request,
        "templates": templates_list,
        "files": files_list
    })
# Request model
class GenerateRequest(BaseModel):
    template: str = Field(..., example="hldaa_template.docx")
    fields: Dict[str, str] = Field(
        ...,
        example={
            "project_name": "MyApp",
            "author_name": "Pramod",
            "date": "2025-09-23"
        }
    )


@app.get("/template-schema", summary="Get template placeholders in ready-to-fill format")
def get_template_schema(template: str = Query(..., example="hld_template.docx")):
    """
    Return template placeholders as a dictionary ready for filling values.
    Example output:
    {
        "template": "hld_template.docx",
        "fields": {
            "system_type": "",
            "author_name": "",
            "date": "",
            "project_name": ""
        }
    }
    """
    template_path = os.path.join(TEMPLATE_DIR, template)
    if not os.path.exists(template_path):
        raise HTTPException(status_code=404, detail=f"Template '{template}' not found")

    try:
        fields = extract_placeholders(template_path)
        field_dict = {field: "" for field in fields}  # keys with empty values
        return {"template": template, "fields": field_dict}
    except Exception as e:
        logging.exception(f"Error extracting schema from template '{template}'")
        raise HTTPException(status_code=500, detail=str(e))



@app.post("/generate", summary="Generate HLD document")
def generate_document(payload: GenerateRequest):
    """
    Accept dynamic field values and generate the final HLD document.
    """
    template_path = os.path.join(TEMPLATE_DIR, payload.template)
    if not os.path.exists(template_path):
        raise HTTPException(status_code=404, detail=f"Template '{payload.template}' not found")

    try:
        expected_fields = extract_placeholders(template_path)
        missing_fields = [f for f in expected_fields if f not in payload.fields]

        if missing_fields:
            raise HTTPException(status_code=400, detail=f"Missing fields: {missing_fields}")

        filename = generate_hld_doc(template_path, payload.fields, output_dir=OUTPUT_DIR)
        download_url = f"/download/{filename}"
        return {"message": "Document generated successfully", "download_url": download_url}

    except HTTPException:
        raise
    except Exception as e:
        logging.exception(f"Error generating document from template '{payload.template}'")
        raise HTTPException(status_code=500, detail=str(e))
################################################################################

################################################################################
from io import BytesIO
@app.post("/preview", summary="Preview HLD document content without saving")
def preview_document(payload: GenerateRequest):
    """
    Fill the template with given fields and return the text content for preview.
    """
    template_path = os.path.join(TEMPLATE_DIR, payload.template)
    if not os.path.exists(template_path):
        raise HTTPException(status_code=404, detail=f"Template '{payload.template}' not found")

    try:
        expected_fields = extract_placeholders(template_path)
        missing_fields = [f for f in expected_fields if f not in payload.fields]

        if missing_fields:
            raise HTTPException(status_code=400, detail=f"Missing fields: {missing_fields}")

        doc = Document(template_path)

        # Replace placeholders in doc in memory
        for para in doc.paragraphs:
            _replace_placeholders_in_paragraph(para, payload.fields)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    _replace_placeholders_in_paragraph(cell, payload.fields)

        # Combine all text for preview
        full_text = "\n".join([p.text for p in doc.paragraphs])
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text += "\n" + cell.text

        return JSONResponse({"preview_text": full_text})

    except Exception as e:
        logging.exception(f"Error generating preview for template '{payload.template}'")
        raise HTTPException(status_code=500, detail=str(e))



@app.get("/download/{filename}", summary="Download generated document")
def download_document(filename: str):
    """
    Serve the generated HLD document.
    """
    file_path = os.path.join(OUTPUT_DIR, filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="Document not found")

    return FileResponse(
        file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename
    )
#
@app.post("/router-config", summary="Send router configuration")
def router_config(payload: dict):
    commands = payload.get("commands")
    if not commands:
        raise HTTPException(status_code=400, detail="No commands provided")

    # Here you can integrate logic to send commands to router via SSH, Netmiko, etc.
    # For example, we'll just echo back the commands for now
    result = f"Commands received:\n{commands}"
    return {"result": result}
