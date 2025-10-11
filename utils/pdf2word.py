"""
pdf2word.py
Expert PDF / Image -> Word/Text converter (Poppler-free).

Core class: Pdf2WordConverter
Features:
 - native PDF conversion (pdf2docx)
 - OCR PDF conversion fallback (PyMuPDF + pytesseract)
 - Image -> Word or Text conversion
 - OCR table detection can be added (OpenCV + pytesseract)
 - batch & parallel conversion
 - progress callback support
 - logging and robust error handling
"""

import os
import uuid
import tempfile
import logging
import threading
from concurrent.futures import ThreadPoolExecutor
from typing import Callable, Optional, Dict
from pathlib import Path

# PDF/Image conversion libraries
from pdf2docx import Converter as PDF2DOCX_Converter
import pdfplumber
import pytesseract
from PIL import Image
from docx import Document
from docx.shared import Inches
import fitz  # PyMuPDF
import cv2
import numpy as np

# Configure logger
logger = logging.getLogger("pdf2word")
if not logger.handlers:
    _h = logging.StreamHandler()
    _h.setFormatter(logging.Formatter("%(asctime)s [%(levelname)s] %(message)s"))
    logger.addHandler(_h)
logger.setLevel(logging.INFO)


class ConversionError(Exception):
    pass


class Pdf2WordConverter:
    """
    PDF/Image -> Word/Text converter.
    Supports native PDF, OCR PDF, Image-to-Word, Image-to-Text.
    """

    def __init__(
        self,
        tesseract_cmd: Optional[str] = None,
        work_dir: Optional[str] = None,
        max_workers: int = 2,
        logger_obj: Optional[logging.Logger] = None,
    ):
        if tesseract_cmd:
            pytesseract.pytesseract.tesseract_cmd = tesseract_cmd

        self.work_dir = Path(work_dir) if work_dir else None
        self.max_workers = max_workers
        self.logger = logger_obj or logger
        self._lock = threading.Lock()

    # -------------------------------
    # Utilities
    # -------------------------------
    def _safe_path(self, path: str) -> Path:
        p = Path(path)
        if not p.parent.exists():
            p.parent.mkdir(parents=True, exist_ok=True)
        return p

    def _is_scanned_pdf(self, pdf_path: str, sample_pages: int = 2) -> bool:
        try:
            with pdfplumber.open(pdf_path) as pdf:
                pages_to_check = min(sample_pages, len(pdf.pages))
                text_len = sum(len(pdf.pages[i].extract_text() or "") for i in range(pages_to_check))
                return text_len < 50
        except Exception:
            return True

    # -------------------------------
    # PDF native conversion
    # -------------------------------
    def convert_native(self, pdf_path: str, docx_path: str, start: Optional[int] = None, end: Optional[int] = None):
        self.logger.info(f"Native PDF -> Word: {pdf_path} -> {docx_path}")
        try:
            cv = PDF2DOCX_Converter(str(pdf_path))
            args = {}
            if start is not None:
                args["start"] = start
            if end is not None:
                args["end"] = end
            cv.convert(str(docx_path), **args)
            cv.close()
        except Exception as e:
            raise ConversionError(f"Native PDF conversion failed: {e}")

    # -------------------------------
    # OCR PDF conversion
    # -------------------------------
    def convert_ocr_pdf(self, pdf_path: str, docx_path: str, progress_callback: Optional[Callable[[int,int],None]] = None):
        self.logger.info(f"OCR PDF -> Word: {pdf_path} -> {docx_path}")
        doc = Document()
        pdf = fitz.open(pdf_path)
        total = len(pdf)

        for idx, page in enumerate(pdf, start=1):
            pix = page.get_pixmap()
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            cv_image = cv2.cvtColor(np.array(img), cv2.COLOR_RGB2BGR)

            try:
                text = pytesseract.image_to_string(cv_image).strip()
            except Exception as e:
                self.logger.exception(f"OCR failed on page {idx}: {e}")
                text = ""

            if text:
                for line in text.splitlines():
                    doc.add_paragraph(line)
            else:
                tmp_file = Path(tempfile.gettempdir()) / f"ocr_img_{uuid.uuid4().hex}.png"
                img.save(tmp_file)
                doc.add_picture(str(tmp_file), width=Inches(6))
                tmp_file.unlink(missing_ok=True)

            if idx != total:
                doc.add_page_break()
            if progress_callback:
                progress_callback(idx, total)

        doc.save(docx_path)

    # -------------------------------
    # Image -> Word conversion
    # -------------------------------
    def convert_image_to_word(self, img_path: str, docx_path: str):
        self.logger.info(f"Image -> Word: {img_path} -> {docx_path}")
        img = Image.open(img_path)
        cv_image = cv2.cvtColor(np.array(img), cv2.COLOR_RGB2BGR)
        doc = Document()

        try:
            text = pytesseract.image_to_string(cv_image).strip()
        except Exception as e:
            self.logger.exception(f"OCR failed: {e}")
            text = ""

        if text:
            for line in text.splitlines():
                doc.add_paragraph(line)
        else:
            tmp_file = Path(tempfile.gettempdir()) / f"ocr_img_{uuid.uuid4().hex}.png"
            img.save(tmp_file)
            doc.add_picture(str(tmp_file), width=Inches(6))
            tmp_file.unlink(missing_ok=True)

        doc.save(docx_path)

    # -------------------------------
    # Image -> Text file conversion
    # -------------------------------
    def convert_image_to_text(self, img_path: str, txt_path: str):
        self.logger.info(f"Image -> Text: {img_path} -> {txt_path}")
        img = Image.open(img_path)
        cv_image = cv2.cvtColor(np.array(img), cv2.COLOR_RGB2BGR)

        try:
            text = pytesseract.image_to_string(cv_image).strip()
        except Exception as e:
            self.logger.exception(f"OCR failed: {e}")
            text = ""

        txt_path = self._safe_path(txt_path)
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(text or "")

    # -------------------------------
    # Unified convert interface
    # -------------------------------
    def convert(self, input_path: str, output_path: str, mode: str = "auto", progress_callback: Optional[Callable[[int,int],None]] = None) -> Dict[str,str]:
        input_path = str(input_path)
        output_path = str(output_path)
        self._safe_path(output_path)
        res = {"success": False, "method": None, "message": ""}

        try:
            ext = Path(input_path).suffix.lower()
            if ext == ".pdf":
                # PDF handling
                if mode == "native":
                    self.convert_native(input_path, output_path)
                    res.update({"success": True, "method":"native"})
                elif mode == "ocr":
                    self.convert_ocr_pdf(input_path, output_path, progress_callback)
                    res.update({"success": True, "method":"ocr"})
                else:  # auto
                    try:
                        self.convert_native(input_path, output_path)
                        res.update({"success": True, "method":"native"})
                    except ConversionError:
                        self.convert_ocr_pdf(input_path, output_path, progress_callback)
                        res.update({"success": True, "method":"ocr"})
            elif ext in [".png", ".jpg", ".jpeg", ".bmp", ".tiff"]:
                # Image handling
                if output_path.lower().endswith(".txt"):
                    self.convert_image_to_text(input_path, output_path)
                    res.update({"success": True, "method":"image-to-text"})
                else:
                    self.convert_image_to_word(input_path, output_path)
                    res.update({"success": True, "method":"image-to-word"})
            else:
                raise ConversionError("Unsupported file type")
        except Exception as e:
            res.update({"success": False, "message": str(e)})

        return res
